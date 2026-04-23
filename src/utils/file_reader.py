"""
Utilitário para leitura de arquivos.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

import fitz
from docx import Document


class LeitorArquivo:
    """Lê e extrai texto de arquivos PDF, DOCX e TXT."""

    FORMATOS_SUPORTADOS = {".pdf", ".docx", ".doc", ".txt"}

    def ler(
        self,
        origem: Union[str, Path, bytes, io.BytesIO],
        nome_arquivo: str = "",
    ) -> str:
        """
        Lê o conteúdo de texto de um arquivo.
        Retorna:
            Texto extraído como string.
        """
        if isinstance(origem, (str, Path)):
            extensao = Path(origem).suffix.lower()
        else:
            extensao = Path(nome_arquivo).suffix.lower() if nome_arquivo else ""

        if extensao == ".pdf":
            return self._ler_pdf(origem)
        elif extensao in (".docx", ".doc"):
            return self._ler_docx(origem)
        elif extensao == ".txt":
            return self._ler_txt(origem)
        else:
            raise ValueError(
                f"Formato '{extensao}' não suportado. "
                f"Use: {', '.join(self.FORMATOS_SUPORTADOS)}"
            )

    def _ler_pdf(self, origem: Union[str, Path, bytes, io.BytesIO]) -> str:
        """Extrai texto de um PDF usando PyMuPDF."""
        try:
            if isinstance(origem, (str, Path)):
                doc = fitz.open(str(origem))
            elif isinstance(origem, bytes):
                doc = fitz.open(stream=origem, filetype="pdf")
            else:
                doc = fitz.open(stream=origem.read(), filetype="pdf")

            paginas = [pagina.get_text() for pagina in doc]
            doc.close()
            return "\n".join(paginas)
        except Exception as e:
            raise IOError(f"Erro ao ler PDF: {e}") from e

    def _ler_docx(self, origem: Union[str, Path, bytes, io.BytesIO]) -> str:
        """Extrai texto de um DOCX usando python-docx."""
        try:
            if isinstance(origem, (str, Path)):
                doc = Document(str(origem))
            elif isinstance(origem, bytes):
                doc = Document(io.BytesIO(origem))
            else:
                doc = Document(origem)

            paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
            # Inclui texto de tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        if celula.text.strip():
                            paragrafos.append(celula.text)
            return "\n".join(paragrafos)
        except Exception as e:
            raise IOError(f"Erro ao ler DOCX: {e}") from e

    def _ler_txt(self, origem: Union[str, Path, bytes, io.BytesIO]) -> str:
        """Lê arquivo de texto simples."""
        try:
            if isinstance(origem, (str, Path)):
                return Path(origem).read_text(encoding="utf-8", errors="ignore")
            elif isinstance(origem, bytes):
                return origem.decode("utf-8", errors="ignore")
            else:
                return origem.read().decode("utf-8", errors="ignore")
        except Exception as e:
            raise IOError(f"Erro ao ler TXT: {e}") from e

    def formato_suportado(self, nome_arquivo: str) -> bool:
        """Verifica se o formato do arquivo é suportado."""
        return Path(nome_arquivo).suffix.lower() in self.FORMATOS_SUPORTADOS

leitor = LeitorArquivo()

def ler_arquivo(
    origem: Union[str, Path, bytes, io.BytesIO],
    nome_arquivo: str = "",
) -> str:
    return leitor.ler(origem, nome_arquivo)
